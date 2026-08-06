import io
import json
import csv
import re
import base64
from typing import List, Dict, Any, Tuple, Optional

try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None


# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def _clean_str(s: str) -> str:
    """Collapse multiple spaces/tabs into one space and strip."""
    if not s:
        return ""
    return re.sub(r'[ \t]{2,}', ' ', s).strip()


def _is_img_line(line: str) -> bool:
    return bool(re.match(r'^\s*<img\b', line, re.I))


# ---------------------------------------------------------------------------
# PHASE 1 — Format-specific text+image extraction
# ---------------------------------------------------------------------------

def _extract_text_from_binary_doc(file_bytes: bytes) -> str:
    """Extract text from legacy binary Word (.doc) files via olefile."""
    text_chunks: List[str] = []
    try:
        import olefile
        if olefile.isOleFile(io.BytesIO(file_bytes)):
            ole = olefile.OleFileIO(io.BytesIO(file_bytes))
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument').read()
                for chunk in re.findall(rb'[\x20-\x7E\t\r\n]{4,}', stream):
                    text_chunks.append(chunk.decode('ascii', errors='ignore'))
    except Exception:
        pass
    if not text_chunks:
        for m in re.findall(rb'(?:[\x20-\x7E][\x00]){4,}', file_bytes):
            try:
                decoded = m.decode('utf-16-le', errors='ignore').strip()
                if len(decoded) > 3:
                    text_chunks.append(decoded)
            except Exception:
                pass
        for m in re.findall(rb'[\x20-\x7E\t\r\n]{4,}', file_bytes):
            decoded = m.decode('ascii', errors='ignore').strip()
            if len(decoded) > 3:
                text_chunks.append(decoded)
    return "\n".join(text_chunks)


def _blob_to_png_b64(blob: bytes) -> Optional[str]:
    """Convert any image byte blob (including WMF/EMF/TIFF/BMP) into a clean browser-renderable PNG base64 data URI."""
    if not blob or len(blob) < 50:
        return None
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(blob))
        if img.mode in ('CMYK', 'P', '1', 'L', 'LA'):
            img = img.convert('RGBA' if 'transparency' in img.info or img.mode == 'LA' else 'RGB')
        out_buf = io.BytesIO()
        img.save(out_buf, format='PNG')
        b64 = base64.b64encode(out_buf.getvalue()).decode('ascii')
        return f'data:image/png;base64,{b64}'
    except Exception:
        try:
            b64 = base64.b64encode(blob).decode('ascii')
            mime = "image/png"
            if blob.startswith(b'\xff\xd8\xff'):
                mime = "image/jpeg"
            elif blob.startswith(b'GIF8'):
                mime = "image/gif"
            return f'data:{mime};base64,{b64}'
        except Exception:
            return None


def _extract_docx_paragraph_text_and_images(p, doc) -> str:
    """
    Extract text + embedded images (DrawingML blip + VML imagedata + shapes + drawings)
    from a python-docx Paragraph. Returns text with inline <img> tags.
    """
    if not p:
        return ""
    parts: List[str] = []
    seen_rIds: set = set()

    def _try_add_image(node, container):
        for attr_key, attr_val in node.attrib.items():
            if attr_key.endswith(('embed', 'id', 'relid', 'href', 'link')) or str(attr_val).startswith('rId'):
                rId = str(attr_val)
                if rId in seen_rIds:
                    continue

                image_part = None
                if doc and hasattr(doc, 'part') and hasattr(doc.part, 'related_parts') and rId in doc.part.related_parts:
                    image_part = doc.part.related_parts[rId]
                elif hasattr(p, 'part') and hasattr(p.part, 'related_parts') and rId in p.part.related_parts:
                    image_part = p.part.related_parts[rId]
                elif doc and hasattr(doc, 'part') and hasattr(doc.part, 'rels') and rId in doc.part.rels:
                    image_part = getattr(doc.part.rels[rId], 'target_part', None)
                elif hasattr(p, 'part') and hasattr(p.part, 'rels') and rId in p.part.rels:
                    image_part = getattr(p.part.rels[rId], 'target_part', None)

                if image_part:
                    seen_rIds.add(rId)
                    blob = getattr(image_part, 'blob', b"")
                    img_data_uri = _blob_to_png_b64(blob)
                    if img_data_uri:
                        container.append(
                            f'<img src="{img_data_uri}" '
                            f'class="q-inline-img" '
                            f'style="max-width:400px;max-height:350px;width:auto;height:auto;'
                            f'object-fit:contain;display:block;margin:8px 0;" />'
                        )

    image_xpath = './/*[local-name()="blip" or local-name()="imagedata" or local-name()="graphicData" or local-name()="shape" or local-name()="pic" or local-name()="inline" or local-name()="anchor"]'

    # Per-run extraction
    for run in p.runs:
        for node in run._r.xpath(image_xpath):
            _try_add_image(node, parts)
        if run.text:
            parts.append(run.text)

    # Paragraph-level fallback (catches images not inside any run)
    for node in p._p.xpath(image_xpath):
        _try_add_image(node, parts)

    if not parts:
        return ""
    if all(p.startswith('<img') for p in parts):
        return "".join(parts)
    return " ".join(parts).strip()


def _extract_docx_cell_text_and_images(cell, doc) -> str:
    cell_parts = []
    for p in cell.paragraphs:
        txt = _extract_docx_paragraph_text_and_images(p, doc)
        if txt:
            cell_parts.append(txt)
    return "\n".join(cell_parts).strip()


def _docx_table_to_html(table, doc) -> str:
    """
    Convert a python-docx Table object into an HTML <table> string.
    Used for tables that appear within a question's content (bar graphs,
    comparison tables, data tables, etc.) so they render properly in the frontend.
    """
    rows_html = []
    for row_idx, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            cell_text = _extract_docx_cell_text_and_images(cell, doc)
            tag = "th" if row_idx == 0 else "td"
            cells_html.append(f"<{tag}>{cell_text}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return (
        '<table class="q-inline-table" style="border-collapse:collapse;margin:8px 0;'
        'font-size:0.875rem;width:auto;max-width:100%;">'
        + "".join(rows_html) +
        "</table>"
    )



def _docx_to_lines(file_bytes: bytes) -> List[str]:
    """
    Convert a DOCX to a flat list of text lines, preserving intra-paragraph
    newlines and extracting embedded images as <img> tags.
    Falls back to _extract_text_from_binary_doc for legacy .doc files.
    """
    if not docx:
        raw = _extract_text_from_binary_doc(file_bytes) or file_bytes.decode("utf-8", errors="ignore")
        return [l for l in raw.splitlines() if l.strip()]

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception:
        raw = _extract_text_from_binary_doc(file_bytes) or file_bytes.decode("utf-8", errors="ignore")
        return [l for l in raw.splitlines() if l.strip()]

    # --- Try structured table extraction first ---
    table_qs = _try_table_extraction(doc)
    if table_qs is not None:
        return table_qs  # Special sentinel: list of dicts, not strings

    # --- Extract body elements in document order ---
    out_lines: List[str] = []
    try:
        for child in doc.element.body:
            if child.tag.endswith('p'):
                p_obj = docx.text.paragraph.Paragraph(child, doc)
                txt = _extract_docx_paragraph_text_and_images(p_obj, doc)
                if txt:
                    # KEY FIX: split on internal newlines so multi-line paragraphs
                    # (question + options + answer in one paragraph) become separate lines
                    for sub in txt.split('\n'):
                        sub = sub.strip()
                        if sub:
                            out_lines.append(sub)
            elif child.tag.endswith('tbl'):
                t = docx.table.Table(child, doc)
                # Check if this looks like a question table (handled by _try_table_extraction)
                # If not, convert to HTML to preserve structure
                headers = []
                try:
                    if t.rows:
                        headers = [cell.text.strip().lower() for cell in t.rows[0].cells]
                except Exception:
                    pass
                q_keywords = ["question", "title", "prompt", "stem", "problem", "item", "query", "statement"]
                is_question_table = any(any(k in h for k in q_keywords) or h in ["q", "q.", "qno", "q.no"] for h in headers)

                if is_question_table:
                    # Let _try_table_extraction handle this (already ran above)
                    for row in t.rows:
                        row_parts = []
                        for cell in row.cells:
                            ct = _extract_docx_cell_text_and_images(cell, doc)
                            if ct:
                                row_parts.append(ct)
                        if row_parts:
                            row_text = " ".join(row_parts)
                            for sub in row_text.split('\n'):
                                sub = sub.strip()
                                if sub:
                                    out_lines.append(sub)
                else:
                    # Inline table (data/chart/comparison) — preserve as HTML
                    html_table = _docx_table_to_html(t, doc)
                    if html_table:
                        out_lines.append(html_table)
    except Exception:
        for p in doc.paragraphs:
            txt = _extract_docx_paragraph_text_and_images(p, doc)
            if txt:
                for sub in txt.split('\n'):
                    sub = sub.strip()
                    if sub:
                        out_lines.append(sub)

    return out_lines


def _try_table_extraction(doc) -> Optional[Any]:
    """
    If the DOCX has structured tables with question headers, extract directly
    into question dicts. Returns None if no suitable tables found.
    """
    table_questions = []
    table_sections = set()

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        q_keywords = ["question", "title", "prompt", "stem", "problem", "item", "query", "statement"]
        has_q = any(any(k in h for k in q_keywords) or h in ["q", "q.", "qno", "q.no"] for h in headers)
        if not has_q or len(table.rows) <= 1:
            continue

        q_col = ans_col = sec_col = marks_col = -1
        opt_cols = []
        for ci, h in enumerate(headers):
            if any(k in h for k in q_keywords) or h in ["q", "q.", "qno", "q.no"]:
                if q_col == -1:
                    q_col = ci
            elif any(k in h for k in ["answer", "correct", "ans", "key"]):
                ans_col = ci
            elif any(k in h for k in ["section", "category", "topic"]):
                sec_col = ci
            elif any(k in h for k in ["mark", "point", "score"]):
                marks_col = ci
            elif any(k in h for k in ["option", "opt", "choice"]) or h in list("abcdef") + ["1","2","3","4"]:
                opt_cols.append(ci)

        for row in table.rows[1:]:
            cells = [_extract_docx_cell_text_and_images(cell, doc) for cell in row.cells]
            if not any(cells):
                continue
            qt = _clean_str(cells[q_col]) if q_col != -1 and q_col < len(cells) else ""
            if not qt:
                continue
            qt = re.sub(r'^(?:q(?:uestion)?\s*[\#\.\-]?\s*\d+[\.\)\:\-]?|\d+[\.\)\:\-]|\(\d+\)|\[\d+\])\s*', '', qt, flags=re.IGNORECASE).strip() or qt
            opts = [_clean_str(cells[oi]) for oi in opt_cols if oi < len(cells) and cells[oi]]
            raw_ans = _clean_str(cells[ans_col]) if ans_col != -1 and ans_col < len(cells) else ""
            sec = _clean_str(cells[sec_col]) if sec_col != -1 and sec_col < len(cells) and cells[sec_col] else "General"
            marks = 1
            if marks_col != -1 and marks_col < len(cells) and cells[marks_col]:
                try:
                    marks = int(float(re.sub(r'[^\d.]', '', cells[marks_col])))
                except (ValueError, TypeError):
                    marks = 1
            mapped = _map_correct_answer(raw_ans, opts)
            qtype = "multiple" if isinstance(mapped, list) and len(mapped) > 1 else ("mcq" if len(opts) >= 2 else "text")
            table_sections.add(sec)
            table_questions.append({
                "id": f"q-{len(table_questions)+1}",
                "type": qtype,
                "question": qt,
                "options": opts,
                "correctAnswer": mapped,
                "section": sec,
                "marks": marks
            })

    if table_questions:
        # Return a special marker so the caller can return directly
        return ("TABLE_RESULT", table_questions, list(table_sections))
    return None


def _pdf_to_lines(file_bytes: bytes) -> List[str]:
    """Extract text lines + inline page images from a PDF using fitz or pypdf."""
    out: List[str] = []

    # 1. Try PyMuPDF (fitz) first for best extraction
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_imgs = []
            for img_info in page.get_images():
                xref = img_info[0]
                try:
                    base_img = doc.extract_image(xref)
                    if base_img and "image" in base_img:
                        img_bytes = base_img["image"]
                        img_data_uri = _blob_to_png_b64(img_bytes)
                        if img_data_uri:
                            page_imgs.append(
                                f'<img src="{img_data_uri}" '
                                f'class="q-inline-img" '
                                f'style="max-width:400px;max-height:350px;width:auto;height:auto;'
                                f'object-fit:contain;display:block;margin:8px 0;" />'
                            )
                except Exception:
                    pass

            page_text = page.get_text("text") or ""
            for l in page_text.splitlines():
                ls = l.strip()
                if not re.match(r'^(?:page\s+\d+(?:\s+of\s+\d+)?|\d+\s*/\s*\d+|---\s*page\s+\d+\s*---)$', ls, re.I) and ls:
                    out.append(ls)
            out.extend(page_imgs)
        if out:
            return out
    except Exception:
        pass

    # 2. Fallback to pypdf
    if pypdf:
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                txt = page.extract_text() or ""
                img_tags = []
                try:
                    if hasattr(page, 'images') and page.images:
                        for img in page.images:
                            data = getattr(img, 'data', None)
                            img_data_uri = _blob_to_png_b64(data)
                            if img_data_uri:
                                img_tags.append(
                                    f'<img src="{img_data_uri}" '
                                    f'class="q-inline-img" '
                                    f'style="max-width:400px;max-height:350px;width:auto;height:auto;'
                                    f'object-fit:contain;display:block;margin:8px 0;" />'
                                )
                except Exception:
                    pass
                for l in txt.splitlines():
                    ls = l.strip()
                    if not re.match(r'^(?:page\s+\d+(?:\s+of\s+\d+)?|\d+\s*/\s*\d+|---\s*page\s+\d+\s*---)$', ls, re.I) and ls:
                        out.append(ls)
                out.extend(img_tags)
            return out
        except Exception:
            pass

    return [l for l in file_bytes.decode("utf-8", errors="ignore").splitlines() if l.strip()]


# ---------------------------------------------------------------------------
# PHASE 2 — Universal state-machine question splitter
# ---------------------------------------------------------------------------

# Compiled regexes (module-level for performance)
_RE_Q_HEADER = re.compile(
    r'^\s*'
    r'(?:'
        r'(?P<prefix>.+?)\s+'              # optional section prefix like "Core ABAP programming"
    r')?'
    r'(?:'
        r'(?:Q(?:uestion)?\s*)'            # "Question " or "Q"
        r'(?P<qnum1>\d+)'                  # question number
        r'[\.\)\:\-]?\s*'
    r'|'
        r'(?P<qnum2>\d+)'                  # bare number
        r'\s*[\.\)\:\-]\s+'                # followed by separator
    r')'
    r'(?P<rest>.*)',
    re.IGNORECASE
)

_RE_LETTERED_OPT = re.compile(
    r'^\s*'
    r'\*?\s*'
    r'(?:'
        r'\(?[A-Fa-f]\)'                   # (A) or A)
        r'|[A-Fa-f]\s*[\.\)\-\:\–]'        # A. A) A- A:
        r'|\[[A-Fa-f]\]'                   # [A]
        r'|\b(?:option|choice)\s+[A-Fa-f0-9]\s*[\.\)\-\:\–]?'  # option A.
        r'|\(?[1-9]\d?\)'                  # (1) or 1)
        r'|\[[1-9]\d?\]'                   # [1]
    r')'
    r'\s*',
    re.IGNORECASE
)

_RE_ANS_LINE = re.compile(
    r'^\s*(?:ans(?:wer)?|correct\s*(?:answer|choice|option)?s?|key|right\s*answer|ans\s*key)\s*[:\-\s]+(.+)',
    re.IGNORECASE
)

_RE_SECTION = re.compile(
    r'^(?:\[?\s*section\b[:\-\s]*|section\s+\d+[:\-\s]*|\[\s*category\b[:\-\s]*|category\s*:\s*|subject\s*:\s*)(.*)',
    re.IGNORECASE
)

_RE_MARKS = re.compile(
    r'\[?\s*(\d+)\s*marks?\s*\]?',
    re.IGNORECASE
)

_RE_ANS_KEY_HEADER = re.compile(
    r'^\s*\[?\s*(?:answer\s*key[s]?|answer\s*sheet|answers)\s*\]?\s*[:\-]?\s*$',
    re.IGNORECASE
)

_RE_ANS_KEY_ITEM = re.compile(
    r'^(?:Q(?:uestion)?\s*)?(\d+)\s*[\.\:\-\s\)]+\s*([A-Za-z0-9,\s&\-\/]+)',
    re.IGNORECASE
)


def _detect_question_header(line: str) -> Optional[Dict]:
    """
    Try to match a question header line. Returns dict with keys:
      prefix, qnum, rest  — or None if not a match.
    """
    if _is_img_line(line):
        return None
    m = _RE_Q_HEADER.match(line)
    if not m:
        return None
    prefix = _clean_str(m.group('prefix') or '')
    qnum_str = m.group('qnum1') or m.group('qnum2')
    rest = _clean_str(m.group('rest') or '')
    if not qnum_str:
        return None
    qnum = int(qnum_str)
    return {"prefix": prefix, "qnum": qnum, "rest": rest}


def _detect_lettered_option(line: str) -> bool:
    """Check if a line starts with a lettered option marker."""
    if _is_img_line(line):
        return False
    return bool(_RE_LETTERED_OPT.match(line))


def _strip_option_prefix(line: str) -> str:
    """Remove the option letter prefix from a line."""
    return _clean_str(_RE_LETTERED_OPT.sub('', line))


def _detect_answer_line(line: str) -> Optional[str]:
    """If line is an answer line, return the answer text. Else None."""
    m = _RE_ANS_LINE.match(line)
    return m.group(1).strip() if m else None


def _extract_marks(text: str) -> Tuple[str, int]:
    """Extract marks from text like '[2 marks]'. Returns (cleaned_text, marks)."""
    m = _RE_MARKS.search(text)
    if m:
        marks = int(m.group(1))
        cleaned = _RE_MARKS.sub('', text).strip()
        return cleaned, marks
    return text, 1


def _detect_structural_options(body_lines: List[str], answer_text: str = "") -> Tuple[List[str], List[str]]:
    """
    Detect unlettered options from body lines using question stem boundary and structural repetition.
    Returns (prompt_lines, option_lines).
    """
    if not body_lines:
        return [], []

    non_img = [l for l in body_lines if not _is_img_line(l)]
    if len(non_img) < 2:
        return body_lines, []

    # Regex indicators for question statements, instructions, and code lines
    q_words_re = re.compile(
        r'^(?:'
        r'what|which|how|why|where|when|who|define|explain|calculate|'
        r'identify|select|choose|find|determine|true\s+or\s+false|'
        r'name\s+the|list\s+the|state\s+the|describe|compare|'
        r'differentiate|distinguish|given|consider|to\s+give|'
        r'in\s+which|for\s+which|is\s+it|can\s+you|does\s+the|'
        r'do\s+the|are\s+the|is\s+the|was\s+the|were\s+the'
        r')\b',
        re.IGNORECASE
    )

    q_indicator_re = re.compile(
        r'(?:'
        r'valid\s+statement|correct\s+statement|correct\s+answer|correct\s+choice|'
        r'correct\s+option|following\s+is\s+true|following\s+are\s+true|'
        r'correct\s+output|select\s+\d+|choose\s+\d+|there\s+are\s+\d+|'
        r'note:|\[select|\(select|\[choose|\(choose'
        r')',
        re.IGNORECASE
    )

    code_indicator_re = re.compile(
        r'^\s*(?:'
        r'DATA\b|TYPES?:|SELECT\b|FROM\b|WHERE\b|def\b|class\b|'
        r'public\b|private\b|function\b|import\b|\#include\b|var\b|'
        r'let\b|const\b|return\b|struct\b|package\b|using\s+namespace|'
        r'if\b|else\b|for\b|while\b|print\(|System\.out'
        r')',
        re.IGNORECASE
    )

    # Search for the LAST line that belongs to the question prompt/stem
    last_q_stmt_idx = -1
    for i, line in enumerate(body_lines):
        if _is_img_line(line):
            continue
        stripped = line.strip()
        is_q_stem_line = (
            '?' in stripped or
            bool(q_words_re.search(stripped)) or
            bool(q_indicator_re.search(stripped)) or
            bool(code_indicator_re.search(stripped))
        )
        if is_q_stem_line:
            last_q_stmt_idx = i

    # Strategy 1: Options must come AFTER the last question stem line
    if last_q_stmt_idx >= 0 and last_q_stmt_idx < len(body_lines) - 1:
        candidate_opts = [l for l in body_lines[last_q_stmt_idx + 1:] if not _is_img_line(l)]
        if len(candidate_opts) >= 2:
            prompt = body_lines[:last_q_stmt_idx + 1]
            for k in range(last_q_stmt_idx + 1, len(body_lines)):
                if _is_img_line(body_lines[k]):
                    prompt.append(body_lines[k])
                else:
                    break
            return prompt, candidate_opts

    # Strategy 2: Answer-text matching
    if answer_text:
        answer_clean = answer_text.strip().lower()
        match_idx = -1
        for i, line in enumerate(non_img):
            if line.strip().lower() == answer_clean or answer_clean in line.strip().lower():
                match_idx = i
                break

        if match_idx >= 0 and match_idx > 0:
            similar_indices = _find_similar_group(non_img, match_idx)
            if len(similar_indices) >= 2:
                min_opt_idx = min(similar_indices)
                prompt = non_img[:min_opt_idx]
                opts = [non_img[i] for i in sorted(similar_indices)]
                if prompt:
                    return prompt, opts

    # Strategy 3: Repeated prefix pattern detection
    prefix_groups = _detect_prefix_groups(non_img)
    if prefix_groups:
        best_group = max(prefix_groups, key=len)
        if len(best_group) >= 2 and len(best_group) <= 20:
            min_idx = min(best_group)
            if min_idx > 0:
                prompt = non_img[:min_idx]
                opts = [non_img[i] for i in sorted(best_group)]
                return prompt, opts

    # Strategy 4: Length-based similarity grouping
    if len(non_img) >= 3:
        for split_at in range(1, len(non_img) - 1):
            candidate_opts = non_img[split_at:]
            if len(candidate_opts) >= 2 and _are_structurally_similar(candidate_opts):
                prompt = non_img[:split_at]
                return prompt, candidate_opts

    # Fallback: answer key letters
    if answer_text:
        letters = re.findall(r'\b[A-Ea-e]\b', answer_text)
        max_idx = max((ord(l.upper()) - ord('A') for l in letters), default=-1) + 1
        num_opts = max(4, max_idx) if letters else 0

        if num_opts > 0 and len(non_img) > num_opts:
            prompt = non_img[:-num_opts]
            opts = non_img[-num_opts:]
            return prompt, opts

    return body_lines, []


def _find_similar_group(lines: List[str], anchor_idx: int) -> List[int]:
    """
    Given an anchor line, find all lines that are structurally similar to it.
    Returns list of indices.
    """
    if anchor_idx >= len(lines):
        return [anchor_idx]

    anchor = lines[anchor_idx].strip()
    anchor_len = len(anchor)
    anchor_words = anchor.split()
    anchor_first_word = anchor_words[0].lower() if anchor_words else ""

    similar = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        words = stripped.split()
        first_word = words[0].lower() if words else ""
        line_len = len(stripped)

        # Check similarity criteria
        is_similar = False

        # Same first word and similar length
        if first_word == anchor_first_word and anchor_len > 0:
            ratio = line_len / anchor_len if anchor_len > 0 else 0
            if 0.3 <= ratio <= 3.0:
                is_similar = True

        # Similar word count and length
        if not is_similar and len(anchor_words) > 0:
            word_ratio = len(words) / len(anchor_words)
            len_ratio = line_len / anchor_len if anchor_len > 0 else 0
            if 0.4 <= word_ratio <= 2.5 and 0.3 <= len_ratio <= 3.0:
                # Check it's a declarative sentence (not a question)
                if not stripped.endswith('?'):
                    is_similar = True

        if is_similar:
            similar.append(i)

    return similar if len(similar) >= 2 else [anchor_idx]


def _detect_prefix_groups(lines: List[str]) -> List[List[int]]:
    """
    Detect groups of lines that share a common prefix pattern (first 1-3 words).
    Returns list of groups, where each group is a list of line indices.
    """
    if len(lines) < 3:
        return []

    # Build prefix map: first word -> list of indices
    prefix_map: Dict[str, List[int]] = {}
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or _is_img_line(stripped):
            continue
        words = stripped.split()
        if not words:
            continue
        # Use first 1-2 words as prefix key
        prefix_key = words[0].lower()
        if prefix_key not in prefix_map:
            prefix_map[prefix_key] = []
        prefix_map[prefix_key].append(i)

    # Filter groups: need at least 3 lines with same prefix
    groups = []
    for prefix, indices in prefix_map.items():
        if len(indices) >= 3:
            # Check they are somewhat contiguous (within a reasonable range)
            min_i, max_i = min(indices), max(indices)
            span = max_i - min_i + 1
            # The group should cover at least 50% of the span
            if len(indices) >= span * 0.4:
                groups.append(indices)

    return groups


def _are_structurally_similar(lines: List[str]) -> bool:
    """
    Check if a group of lines are structurally similar (likely options).
    Criteria: similar length, similar word count, similar structure.
    """
    if len(lines) < 2:
        return False

    cleaned = [l.strip() for l in lines if l.strip() and not _is_img_line(l)]
    if len(cleaned) < 2:
        return False

    lengths = [len(l) for l in cleaned]
    word_counts = [len(l.split()) for l in cleaned]
    avg_len = sum(lengths) / len(lengths)
    avg_words = sum(word_counts) / len(word_counts)

    if avg_len < 5:  # Too short to be meaningful options
        return False

    # Check length variance
    if avg_len > 0:
        len_variance = sum((l - avg_len) ** 2 for l in lengths) / len(lengths)
        len_cv = (len_variance ** 0.5) / avg_len  # coefficient of variation
        if len_cv < 0.8:  # Similar lengths
            return True

    # Check if they share common first words
    first_words = [l.split()[0].lower() if l.split() else "" for l in cleaned]
    most_common_first = max(set(first_words), key=first_words.count)
    if first_words.count(most_common_first) >= len(cleaned) * 0.6:
        return True

    # Check similar word counts
    if avg_words > 0:
        word_variance = sum((w - avg_words) ** 2 for w in word_counts) / len(word_counts)
        word_cv = (word_variance ** 0.5) / avg_words
        if word_cv < 0.5:
            return True

    return False


def _parse_lines(lines: List[str]) -> Tuple[List[Dict[str, Any]], List[str]]:
    """
    Universal two-pass block-based parser that converts a flat list of text lines
    into structured question dicts.

    Pass 1: Extract answer keys, collect question blocks.
    Pass 2: Analyze each block to separate question text from options.

    Handles:
    - Numbered questions: `1.`, `Q1.`, `Question 1`, `Topic Question 1`
    - Lettered options: `A)`, `a.`, `(A)`, `[A]`, `A -`
    - UNLETTERED options: detected via structural repetition patterns
    - Answer lines: `Answer: A`, `Correct Answer: c) $TMP`, `Answer A C E`
    - Separate answer key sections
    - Section headers
    - Inline images and HTML tables
    - Marks extraction
    - Unnumbered questions (fallback)
    """

    # ---- Pass 1: Extract answer key if present ----
    answer_key: Dict[int, str] = {}
    content_lines: List[str] = []
    in_answer_key = False

    for line in lines:
        cleaned = _clean_str(line)
        if not cleaned:
            continue
        if _RE_ANS_KEY_HEADER.match(cleaned):
            in_answer_key = True
            continue
        if in_answer_key:
            m = _RE_ANS_KEY_ITEM.match(cleaned)
            if m:
                answer_key[int(m.group(1))] = m.group(2).strip()
                continue
            # Check for inline compact keys like "1 A  2 B  3 C"
            inline = re.findall(r'(\d+)\s*[\:\-\.\s]+([A-Za-z0-9]+)', cleaned)
            if inline:
                for qk, av in inline:
                    try:
                        answer_key[int(qk)] = av
                    except ValueError:
                        pass
                continue
            # If line doesn't match answer key pattern, answer key section ended
            in_answer_key = False
        content_lines.append(cleaned)

    # ---- Pass 2: Collect question blocks ----
    # A "block" is all lines between one question header and the next.
    blocks: List[Dict[str, Any]] = []
    current_block: Optional[Dict[str, Any]] = None
    current_section = "General"

    for idx, line in enumerate(content_lines):
        # Lettered or Numbered Option? (Check before question header if current_block is active)
        if current_block is not None and _detect_lettered_option(line):
            # Check for inline options e.g. "(A) 3  (B) 4  (C) 5  (D) 6"
            split_parts = [p.strip() for p in re.split(r'\s+(?=\(?\*?\s*(?:[A-Za-z0-9][\.\)\:\-]|\[[A-Za-z0-9]\]|\([A-Za-z0-9]\))\s*)', line) if p.strip()]
            if len(split_parts) >= 2:
                for part in split_parts:
                    if part.lstrip().startswith('*'):
                        current_block["answer"] = _strip_option_prefix(part)
                    current_block["lettered_options"].append(part)
            else:
                if line.lstrip().startswith('*'):
                    current_block["answer"] = _strip_option_prefix(line)
                current_block["lettered_options"].append(line)
            continue

        # Section header?
        sec_m = _RE_SECTION.match(line)
        if sec_m:
            val = _clean_str(sec_m.group(1).strip("] "))
            if val:
                if current_block:
                    blocks.append(current_block)
                    current_block = None
                current_section = val
                continue

        # Question header?
        qh = _detect_question_header(line)
        if qh:
            if current_block:
                blocks.append(current_block)
            prefix = qh["prefix"]
            if prefix and len(prefix) > 2:
                if not re.match(r'^\d+[\.\)\-\s]*$', prefix):
                    current_section = prefix
            current_block = {
                "qnum": qh["qnum"],
                "section": current_section,
                "body": [],
                "lettered_options": [],
                "answer": "",
                "marks": 1,
                "pending_imgs": [],
            }
            if qh["rest"]:
                current_block["body"].append(qh["rest"])
            continue

        # Answer line?
        ans_text = _detect_answer_line(line)
        if ans_text is not None:
            if current_block:
                current_block["answer"] = ans_text
            continue

        # Image line?
        if _is_img_line(line):
            if current_block:
                if current_block["answer"]:
                    current_block["pending_imgs"].append(line)
                else:
                    current_block["body"].append(line)
            continue

        # Unnumbered question detection (fallback)
        if current_block is None:
            next_line = content_lines[idx + 1] if idx + 1 < len(content_lines) else ""
            is_q_phrase = bool(re.match(
                r'^(?:what|which|how|why|where|when|who|define|explain|calculate|'
                r'identify|select|find|true\s+or\s+false|name\s+the|list\s+the|'
                r'state\s+the|describe|compare|differentiate|distinguish|'
                r'to\s+give|in\s+which|for\s+which|is\s+it|can\s+you|'
                r'does\s+the|do\s+the|are\s+the|is\s+the|was\s+the|were\s+the)\b',
                line, re.I
            ))
            next_is_opt = _detect_lettered_option(next_line)
            next_is_ans = _detect_answer_line(next_line) is not None

            if "?" in line or (is_q_phrase and (next_is_opt or next_is_ans)):
                current_block = {
                    "qnum": len(blocks) + 1,
                    "section": current_section,
                    "body": [line],
                    "lettered_options": [],
                    "answer": "",
                    "marks": 1,
                    "pending_imgs": [],
                }
                continue

        # Default: append to current block body
        if current_block is not None:
            if not current_block["answer"]:
                current_block["body"].append(line)

    if current_block:
        blocks.append(current_block)

    # ---- Pass 3: Analyze each block → produce question dicts ----
    questions: List[Dict[str, Any]] = []
    sections_set: set = set()

    for block in blocks:
        q_num = block["qnum"]
        q_section = block["section"]
        q_body = block["body"]
        q_lettered = block["lettered_options"]
        q_answer_raw = block["answer"]
        q_marks = block["marks"]

        if not q_body and not q_lettered:
            continue

        # Build prompt and options
        prompt_parts: List[str] = []
        options: List[str] = []

        if q_lettered:
            # Explicit lettered options — use them directly
            prompt_parts = list(q_body)
            for opt_line in q_lettered:
                cleaned = _strip_option_prefix(opt_line) or _clean_str(opt_line)
                if cleaned:
                    options.append(cleaned)
        else:
            # No lettered options — use structural detection
            detected_prompt, detected_opts = _detect_structural_options(
                q_body, q_answer_raw
            )
            prompt_parts = detected_prompt
            options = [_clean_str(o) for o in detected_opts if _clean_str(o)]

        if not prompt_parts and options:
            prompt_parts = [options.pop(0)]

        prompt = "\n".join(prompt_parts).strip()
        if not prompt:
            continue

        # Clean question number prefix from prompt
        prompt_clean = _clean_str(re.sub(
            r'^(?:Q(?:uestion)?\s*[\#\.\-]?\s*\d+[\.\)\:\-]?|\d+[\.\)\:\-]|\(\d+\)|\[\d+\])\s*',
            '', prompt, flags=re.IGNORECASE
        )) or prompt

        # Extract marks from prompt
        prompt_clean, detected_marks = _extract_marks(prompt_clean)
        if detected_marks > 1:
            q_marks = detected_marks

        # Resolve answer
        raw_ans = q_answer_raw.strip()
        key_ans = answer_key.get(q_num or (len(questions) + 1))
        if not raw_ans and key_ans:
            raw_ans = str(key_ans)

        # Deduplicate options
        unique_opts: List[str] = []
        seen_opts: set = set()
        for o in options:
            if o and o not in seen_opts:
                seen_opts.add(o)
                unique_opts.append(o)
        options = unique_opts

        mapped_ans: Any = ""
        if raw_ans:
            mapped_ans = _map_correct_answer(raw_ans, options)
        elif options:
            mapped_ans = options[0]

        # Determine type
        q_type = "text"
        if isinstance(mapped_ans, list) and len(mapped_ans) > 1:
            q_type = "multiple"
        elif len(options) >= 2:
            q_type = "mcq"

        # Extract imageUrl
        img_m = re.search(r'<img\s+[^>]*src=["\']([^"\']+)["\']', prompt_clean)
        image_url = img_m.group(1) if img_m else None

        sec = _clean_str(q_section) or "General"
        sections_set.add(sec)

        questions.append({
            "id": f"q-{len(questions) + 1}",
            "type": q_type,
            "question": prompt_clean,
            "options": options,
            "correctAnswer": mapped_ans,
            "section": sec,
            "marks": q_marks,
            "imageUrl": image_url
        })

    return questions, list(sections_set) if sections_set else ["General"]


# ---------------------------------------------------------------------------
# PHASE 3 — Public API
# ---------------------------------------------------------------------------

def parse_questions_file(file_bytes: bytes, filename: str) -> Tuple[List[Dict[str, Any]], List[str]]:
    """
    Universal parser for uploaded test files (.json, .csv, .txt, .docx, .doc, .pdf).
    Supports embedded images, unnumbered questions, arbitrary templates, answer keys.
    Returns: (questions, sections)
    """
    fname = filename.lower()

    if fname.endswith(".json"):
        return _parse_json(file_bytes)
    if fname.endswith(".csv"):
        return _parse_csv(file_bytes)
    if fname.endswith(".docx") or fname.endswith(".doc"):
        result = _docx_to_lines(file_bytes)
        # Check if table extraction returned directly
        if result and isinstance(result, tuple) and result[0] == "TABLE_RESULT":
            return result[1], result[2]
        return _parse_lines(result)
    if fname.endswith(".pdf"):
        return _parse_lines(_pdf_to_lines(file_bytes))

    # Fallback: treat as plain text
    raw = file_bytes.decode("utf-8", errors="replace")
    lines = [_clean_str(l) for l in raw.splitlines() if _clean_str(l)]
    return _parse_lines(lines)


# ---------------------------------------------------------------------------
# JSON parser
# ---------------------------------------------------------------------------

def _parse_json(file_bytes: bytes) -> Tuple[List[Dict[str, Any]], List[str]]:
    text = file_bytes.decode("utf-8", errors="replace")
    try:
        data = json.loads(text)
    except Exception:
        return [], ["General"]

    items = []
    if isinstance(data, list):
        items = data
    elif isinstance(data, dict):
        for key in ["questions", "items", "data", "questionList"]:
            if data.get(key):
                items = data[key]
                break

    questions = []
    sections_set: set = set()

    for idx, item in enumerate(items, 1):
        if not isinstance(item, dict):
            continue
        q_text = _clean_str(str(
            item.get("question") or item.get("questionText") or item.get("title") or
            item.get("prompt") or item.get("stem") or item.get("q") or ""
        ))
        if not q_text:
            continue
        q_text = re.sub(
            r'^(?:q(?:uestion)?\s*[\#\.\-]?\s*\d+[\.\)\:\-]?|\d+[\.\)\:\-]|\(\d+\)|\[\d+\])\s*',
            '', q_text, flags=re.IGNORECASE
        ).strip() or q_text

        raw_options = item.get("options") or item.get("choices") or item.get("opt") or []
        if isinstance(raw_options, str):
            opts = [_clean_str(o) for o in re.split(r'[|;\n]', raw_options) if o.strip()]
        else:
            opts = [_clean_str(str(o)) for o in raw_options if str(o).strip()]
        # Deduplicate
        seen: set = set()
        unique: List[str] = []
        for o in opts:
            if o not in seen:
                seen.add(o)
                unique.append(o)
        opts = unique

        raw_ans = (
            item.get("correctAnswer") or item.get("correct_answer") or item.get("answer") or
            item.get("correct") or item.get("correctAnswers") or item.get("ans") or item.get("key") or ""
        )
        mapped = _map_correct_answer(raw_ans, opts)

        q_type = item.get("type")
        if not q_type:
            if isinstance(mapped, list) and len(mapped) > 1:
                q_type = "multiple"
            elif len(opts) >= 2:
                q_type = "mcq"
            else:
                q_type = "text"

        section = _clean_str(str(item.get("section") or item.get("category") or "General")) or "General"
        sections_set.add(section)

        marks = 1
        try:
            marks = int(float(str(item.get("marks") or item.get("points") or item.get("score") or 1).strip()))
        except (ValueError, TypeError):
            marks = 1

        questions.append({
            "id": f"q-{idx}", "type": q_type, "question": q_text,
            "options": opts, "correctAnswer": mapped, "section": section, "marks": marks
        })

    return questions, list(sections_set) if sections_set else ["General"]


# ---------------------------------------------------------------------------
# CSV parser
# ---------------------------------------------------------------------------

def _parse_csv(file_bytes: bytes) -> Tuple[List[Dict[str, Any]], List[str]]:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [r for r in reader if any(cell.strip() for cell in r)]
    if not rows:
        return [], ["General"]

    header = [c.strip().lower() for c in rows[0]]
    q_kw = ["question", "title", "prompt", "stem", "problem", "item", "statement"]
    has_header = any(any(k in h for k in q_kw) or h in ["q", "q.", "qno", "q.no"] for h in header)

    questions = []
    sections_set: set = set()
    data_rows = rows[1:] if has_header else rows

    for idx, row in enumerate(data_rows, 1):
        if not row:
            continue
        q_text = ""
        opts_list: List[str] = []
        raw_ans = ""
        section = "General"
        marks = 1
        q_type = None

        if has_header:
            rd = {header[i]: row[i].strip() for i in range(min(len(header), len(row)))}
            for k, v in rd.items():
                if any(w in k for w in q_kw) or k in ["q", "q.", "qno", "q.no"]:
                    if not q_text:
                        q_text = _clean_str(v)
                elif any(w in k for w in ["section", "category", "topic"]):
                    if section == "General" and v:
                        section = _clean_str(v)
                elif any(w in k for w in ["answer", "correct", "ans", "key"]):
                    if not raw_ans and v:
                        raw_ans = _clean_str(v)
                elif any(w in k for w in ["mark", "point", "score"]):
                    try:
                        marks = int(float(v))
                    except (ValueError, TypeError):
                        pass
                elif k == "type":
                    q_type = v
            if "options" in rd and rd["options"]:
                opts_list = [_clean_str(o) for o in re.split(r'[|;\n]', rd["options"]) if o.strip()]
            else:
                for k in header:
                    if (any(w in k for w in ["option", "opt", "choice"]) or k in list("abcdef") + ["1","2","3","4"]) \
                            and k not in ["options", "option", "choice", "choices"]:
                        if rd.get(k):
                            opts_list.append(_clean_str(rd[k]))
        else:
            q_text = _clean_str(row[0])
            if len(row) > 2:
                raw_ans = _clean_str(row[-1])
                opts_list = [_clean_str(c) for c in row[1:-1] if c.strip()]

        if not q_text:
            continue
        # Deduplicate
        seen: set = set()
        unique: List[str] = []
        for o in opts_list:
            if o not in seen:
                seen.add(o)
                unique.append(o)
        opts_list = unique
        q_text = re.sub(r'^(?:q(?:uestion)?\s*[\#\.\-]?\s*\d+[\.\)\:\-]?|\d+[\.\)\:\-]|\(\d+\)|\[\d+\])\s*', '', q_text, flags=re.IGNORECASE).strip() or q_text
        sections_set.add(section)
        mapped = _map_correct_answer(raw_ans, opts_list)
        if not q_type:
            if isinstance(mapped, list) and len(mapped) > 1:
                q_type = "multiple"
            elif len(opts_list) >= 2:
                q_type = "mcq"
            else:
                q_type = "text"
        questions.append({
            "id": f"q-{idx}", "type": q_type, "question": q_text,
            "options": opts_list, "correctAnswer": mapped, "section": section, "marks": marks
        })

    return questions, list(sections_set) if sections_set else ["General"]


# ---------------------------------------------------------------------------
# Answer mapping helper
# ---------------------------------------------------------------------------

def _map_correct_answer(raw_ans: Any, options: List[str]) -> Any:
    """Map a raw answer value to the actual option text(s)."""
    if isinstance(raw_ans, list):
        mapped = []
        for a in raw_ans:
            m = _map_correct_answer(a, options)
            if isinstance(m, list):
                mapped.extend(m)
            elif m and m not in mapped:
                mapped.append(m)
        return mapped

    raw_str = _clean_str(str(raw_ans))
    if not raw_str:
        return ""
    if not options:
        return raw_str

    letter_map: Dict[str, int] = {}
    for i in range(26):
        letter_map[chr(ord('A') + i)] = i
        letter_map[chr(ord('a') + i)] = i
    for i in range(1, 100):
        letter_map[str(i)] = i - 1

    clean = raw_str.strip()

    # Exact match first (e.g. option text is literally "Paris" or "4")
    for opt in options:
        if opt.strip().lower() == clean.lower():
            return opt

    # Check if raw_str is a multi-letter/number answer like "B, C, E", "B C E", "A B C", "1 3 5", "A,C,E", "A/B/C"
    parts = [p.strip() for p in re.split(r'[,&/;\-\s]|\band\b', clean, flags=re.IGNORECASE) if p.strip()]
    clean_parts = []
    for p in parts:
        p_clean = re.sub(r'^(?:option|choice|ans(?:wer)?|key|correct)\s*', '', p, flags=re.I).strip()
        if p_clean and p_clean in letter_map:
            clean_parts.append(p_clean)

    if len(clean_parts) >= 2:
        mapped_list = []
        for p in clean_parts:
            idx = letter_map[p]
            if 0 <= idx < len(options):
                opt_text = options[idx]
                if opt_text not in mapped_list:
                    mapped_list.append(opt_text)
        if mapped_list:
            return mapped_list if len(mapped_list) > 1 else mapped_list[0]

    # Single letter/number index match (e.g. "B" or "2")
    lm = re.search(r'^(?:option|choice|ans(?:wer)?|key|\()?([A-Za-z]|\d+)[\)\.]?$', clean, re.I)
    if lm:
        let = lm.group(1)
        if let in letter_map:
            idx = letter_map[let]
            if 0 <= idx < len(options):
                return options[idx]

    # Strip prefix and try again
    stripped = re.sub(r'^(?:option|choice|ans(?:wer)?|\()?([A-Za-z]|\d+)[\)\.\:\s]+', '', clean, flags=re.I).strip()
    if stripped:
        for opt in options:
            if opt.strip().lower() == stripped.lower():
                return opt

    # Substring match
    for opt in options:
        if clean.lower() in opt.strip().lower() or opt.strip().lower() in clean.lower():
            return opt

    return raw_str
